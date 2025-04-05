"""
Text Extraction Utilities for AI Scientific Research Summarizer
Provides high-quality text extraction from various document formats
"""

import os
import sys
import json
import logging
import tempfile
from typing import Dict, Any, Optional, List, Union
from urllib.parse import urlparse

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define supported file types
SUPPORTED_FILE_TYPES = {
    'pdf': ['pdf'],
    'image': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp'],
    'document': ['doc', 'docx', 'txt', 'rtf', 'odt'],
    'spreadsheet': ['csv', 'xls', 'xlsx'],
    'presentation': ['ppt', 'pptx'],
    'webpage': ['html', 'htm', 'url']
}

# Import specialized libraries for each file type
try:
    # PDF processing
    import PyPDF2
    import pdfplumber
    from pdf2image import convert_from_path
    import pytesseract
    HAS_PDF_LIBS = True
except ImportError:
    logger.warning("PDF processing libraries not installed. Run: pip install PyPDF2 pdfplumber pdf2image pytesseract")
    HAS_PDF_LIBS = False

try:
    # Document processing
    import docx
    from odf import text, teletype
    from odf.opendocument import load
    import win32com.client
    HAS_DOC_LIBS = True
except ImportError:
    logger.warning("Document processing libraries not installed. Run: pip install python-docx odfpy pywin32")
    HAS_DOC_LIBS = False

try:
    # Image processing
    from PIL import Image, ImageFilter, ImageEnhance
    import pytesseract
    HAS_IMAGE_LIBS = True
except ImportError:
    logger.warning("Image processing libraries not installed. Run: pip install pillow pytesseract")
    HAS_IMAGE_LIBS = False

try:
    # Spreadsheet processing
    import pandas as pd
    import openpyxl
    HAS_SPREADSHEET_LIBS = True
except ImportError:
    logger.warning("Spreadsheet processing libraries not installed. Run: pip install pandas openpyxl")
    HAS_SPREADSHEET_LIBS = False

try:
    # Web content processing
    import requests
    from bs4 import BeautifulSoup
    import html2text
    HAS_WEB_LIBS = True
except ImportError:
    logger.warning("Web content processing libraries not installed. Run: pip install requests beautifulsoup4 html2text")
    HAS_WEB_LIBS = False

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using multiple methods for best results
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    if not HAS_PDF_LIBS:
        raise ImportError("PDF processing libraries not installed")
    
    text = ""
    
    # Method 1: PyPDF2 for basic text extraction
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page_text = reader.pages[page_num].extract_text() or ""
                text += page_text + "\n\n"
        
        logger.info(f"Extracted {len(text.split())} words using PyPDF2")
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # Method 2: If PyPDF2 didn't extract much text, try pdfplumber
    if len(text.strip()) < 100:
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
            
            logger.info(f"Extracted {len(text.split())} words using pdfplumber")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Method 3: If text is still minimal, try OCR as a last resort
    if len(text.strip()) < 100:
        try:
            text = ""
            images = convert_from_path(file_path)
            for i, image in enumerate(images):
                # Enhance image for better OCR
                image = image.convert('L')  # Convert to grayscale
                image = ImageEnhance.Contrast(image).enhance(2.0)  # Increase contrast
                image = image.filter(ImageFilter.SHARPEN)  # Sharpen image
                
                # Extract text with OCR
                page_text = pytesseract.image_to_string(image) or ""
                text += page_text + "\n\n"
            
            logger.info(f"Extracted {len(text.split())} words using OCR")
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
    
    return text

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from DOC files using win32com
    
    Args:
        file_path: Path to the DOC file
        
    Returns:
        Extracted text from the DOC
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        # Convert relative path to absolute path
        file_path = os.path.abspath(file_path)
        
        # Initialize Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            # Open the document
            doc = word.Documents.Open(file_path)
            
            # Extract text
            text = doc.Content.Text
            
            # Close the document
            doc.Close()
            
            logger.info(f"Extracted {len(text.split())} words from DOC")
            return text
            
        finally:
            # Always quit Word application
            word.Quit()
    
    except Exception as e:
        logger.error(f"DOC extraction failed: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        logger.info(f"Extracted {len(text.split())} words from DOCX")
        return text
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_odt(file_path: str) -> str:
    """
    Extract text from ODT files
    
    Args:
        file_path: Path to the ODT file
        
    Returns:
        Extracted text from the ODT
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        textdoc = load(file_path)
        allparas = textdoc.getElementsByType(text.P)
        text_content = "\n\n".join([teletype.extractText(para) for para in allparas])
        
        logger.info(f"Extracted {len(text_content.split())} words from ODT")
        return text_content
    except Exception as e:
        logger.error(f"ODT extraction failed: {e}")
        return ""

def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from images using OCR
    Supports both local files and URLs
    
    Args:
        file_path: Path to the image file or URL
        
    Returns:
        Extracted text from the image
    """
    if not HAS_IMAGE_LIBS:
        raise ImportError("Image processing libraries not installed")
    
    try:
        logger.info(f"Processing image: {file_path}")
        
        # Open the image
        image = Image.open(file_path)
        
        # Preprocess image for better OCR results
        # 1. Convert to grayscale
        image = image.convert('L')
        
        # 2. Increase contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        
        # 3. Apply sharpening filter
        image = image.filter(ImageFilter.SHARPEN)
        
        # 4. Resize if too small
        if image.width < 1000 or image.height < 1000:
            scale_factor = max(1000 / image.width, 1000 / image.height)
            new_width = int(image.width * scale_factor)
            new_height = int(image.height * scale_factor)
            image = image.resize((new_width, new_height), Image.LANCZOS)
        
        # Extract text using OCR
        text = pytesseract.image_to_string(image)
        
        # Log extraction statistics
        word_count = len(text.split())
        logger.info(f"Extracted {word_count} words from image using OCR")
        
        # If text is very short, try different OCR configurations
        if word_count < 10:
            logger.info("Trying alternative OCR configuration")
            custom_config = r'--oem 1 --psm 3'
            text = pytesseract.image_to_string(image, config=custom_config)
            logger.info(f"Extracted {len(text.split())} words with alternative OCR")
        
        return text
    
    except Exception as e:
        logger.error(f"Image extraction failed: {e}")
        
        # Provide a more helpful error message but don't fail completely
        return f"Image text extraction attempted but encountered issues: {str(e)}. The image may contain little text or be of low quality."

def extract_text_from_excel(file_path: str) -> str:
    """
    Extract text from Excel files
    
    Args:
        file_path: Path to the Excel file
        
    Returns:
        Extracted text from the Excel file
    """
    if not HAS_SPREADSHEET_LIBS:
        raise ImportError("Spreadsheet processing libraries not installed")
    
    try:
        # Read Excel file
        df = pd.read_excel(file_path, sheet_name=None)
        
        text_parts = []
        
        # Process each sheet
        for sheet_name, sheet_df in df.items():
            text_parts.append(f"Sheet: {sheet_name}")
            text_parts.append(sheet_df.to_string(index=False))
            text_parts.append("\n")
        
        text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(text.split())} words from Excel")
        return text
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return ""

def extract_text_from_csv(file_path: str) -> str:
    """
    Extract text from CSV files
    
    Args:
        file_path: Path to the CSV file
        
    Returns:
        Extracted text from the CSV
    """
    if not HAS_SPREADSHEET_LIBS:
        raise ImportError("Spreadsheet processing libraries not installed")
    
    try:
        # Read CSV file
        df = pd.read_csv(file_path)
        
        # Convert to string representation
        text = df.to_string(index=False)
        
        logger.info(f"Extracted {len(text.split())} words from CSV")
        return text
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_html(file_path: str) -> str:
    """
    Extract text from HTML files
    
    Args:
        file_path: Path to the HTML file
        
    Returns:
        Extracted text from the HTML
    """
    if not HAS_WEB_LIBS:
        raise ImportError("Web content processing libraries not installed")
    
    try:
        # Read HTML file
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text()
        
        # Break into lines and remove leading and trailing space
        lines = (line.strip() for line in text.splitlines())
        
        # Break multi-headlines into a line each
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        
        # Drop blank lines
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Extracted {len(text.split())} words from HTML")
        return text
    except Exception as e:
        logger.error(f"HTML extraction failed: {e}")
        return ""

def extract_text_from_url(url: str) -> str:
    """
    Extract text from a URL
    
    Args:
        url: URL to extract text from
        
    Returns:
        Extracted text from the URL
    """
    if not HAS_WEB_LIBS:
        raise ImportError("Web content processing libraries not installed")
    
    try:
        # Fetch URL content
        response = requests.get(url)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text()
        
        # Process text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Extracted {len(text.split())} words from URL")
        return text
    except Exception as e:
        logger.error(f"URL extraction failed: {e}")
        return ""

def preprocess_text(text: str) -> str:
    """
    Preprocess extracted text to improve quality
    
    Args:
        text: Raw extracted text
        
    Returns:
        Preprocessed text
    """
    if not text:
        return ""
    
    # Remove excessive newlines
    text = '\n'.join([line for line in text.splitlines() if line.strip()])
    
    # Remove excessive spaces
    text = ' '.join([word for word in text.split() if word])
    
    return text

def extract_text_from_file(file_path: str, file_type: str = None) -> str:
    """
    Main function to extract text from any file type
    
    Args:
        file_path: Path to the file
        file_type: Type of file (pdf, image, etc.)
        
    Returns:
        Extracted text from the file
    """
    # Determine file type if not provided
    if not file_type:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip('.')
    
    # Normalize file type
    file_type = file_type.lower()
    
    # Extract text based on file type
    try:
        if file_type in SUPPORTED_FILE_TYPES['pdf'] or file_type == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif any(file_type in ext_list for ext_list in [SUPPORTED_FILE_TYPES['image']]) or file_type == 'image':
            text = extract_text_from_image(file_path)
        elif file_type == 'docx':
            text = extract_text_from_docx(file_path)
        elif file_type == 'doc':
            text = extract_text_from_doc(file_path)
        elif file_type == 'odt':
            text = extract_text_from_odt(file_path)
        elif file_type in ['xlsx', 'xls']:
            text = extract_text_from_excel(file_path)
        elif file_type in ['csv']:
            text = extract_text_from_csv(file_path)
        elif file_type in ['html', 'htm']:
            text = extract_text_from_html(file_path)
        elif file_type.startswith('http'):
            text = extract_text_from_url(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Preprocess the extracted text
        text = preprocess_text(text)
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_type} file: {e}", exc_info=True)
        
        # For image files, provide a graceful fallback rather than failing
        if file_type == 'image' or file_type in SUPPORTED_FILE_TYPES['image']:
            return f"Image text extraction attempted but encountered technical issues: {str(e)}. The system will try to process this image with alternative methods."
        
        raise ValueError(f"Failed to extract text from {file_type} file: {str(e)}")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Extract text from various file formats')
    parser.add_argument('file_path', help='Path to the file or URL')
    parser.add_argument('--file-type', help='File type (pdf, image, docx, etc.)')
    parser.add_argument('--output', help='Output file path (optional)')
    
    args = parser.parse_args()
    
    try:
        # Extract text
        text = extract_text_from_file(args.file_path, args.file_type)
        
        # Save or print output
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Text extracted and saved to {args.output}")
        else:
            print(text)
        
        # Return success status
        print(json.dumps({
            "status": "success",
            "text_length": len(text),
            "word_count": len(text.split())
        }))
        
    except Exception as e:
        # Return error status
        print(json.dumps({
            "status": "error",
            "error": str(e)
        }), file=sys.stderr)
        sys.exit(1)
